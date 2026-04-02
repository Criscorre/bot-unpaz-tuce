# herramientas.py
import os
import io
import time
import base64
import tempfile
import requests
from telebot import types

# ─────────────────────────────────────────────
#  ESTADO CONVERSACIONAL
# ─────────────────────────────────────────────
estados_herramientas = {}

# Materias disponibles para el banco
MATERIAS_BANCO = [
    # 1° Año
    "Tecnología y Sociedad",
    "Inglés I",
    "Inglés II",
    "Principios de Economía",
    "Comunicación Institucional",
    "Internet: Infraestructura y redes",
    "Semántica de las interfaces",
    "Introducción al comercio electrónico",
    "Usabilidad, seguridad y Estándares Web",
    # 2° Año
    "Investigación de mercado",
    "Marco legal de negocios electrónicos",
    "Gestión del conocimiento",
    "Desarrollo Web",
    "Formulación, incubación y evaluación de proyectos",
    "Métricas del mundo digital",
    "Desarrollo de Productos y Servicios",
    "Taller de Comunicación",
    "Desarrollos para Dispositivos móviles",
    # 3° Año
    "Calidad y Servicio al Cliente",
    "Marketing digital",
    "Taller de Práctica Integradora",
    "Competencias emprendedoras",
    "Gestión de Proyectos",
]

TIPOS_MATERIAL = [
    "📝 Parcial/Examen",
    "📖 Resumen",
    "📋 Apunte de clase",
    "🔖 Guía de estudio",
    "💡 Trabajo práctico",
]


# ─────────────────────────────────────────────
#  MENÚ PRINCIPAL HERRAMIENTAS
# ─────────────────────────────────────────────

def menu_herramientas(bot, message):
    markup = types.InlineKeyboardMarkup(row_width=1)
    markup.add(
        types.InlineKeyboardButton("📂 Banco de material",        callback_data="her_banco"),
        types.InlineKeyboardButton("📸 Foto → Word / PDF",        callback_data="her_ocr"),
    )
    bot.send_message(
        message.chat.id,
        "🛠️ *Herramientas Estudiantes*\n\n"
        "Todo lo que necesitás para estudiar mejor.\n"
        "¿Qué querés usar?",
        reply_markup=markup,
        parse_mode="Markdown"
    )


# ─────────────────────────────────────────────
#  BANCO DE MATERIAL
# ─────────────────────────────────────────────

def menu_banco(bot, call):
    markup = types.InlineKeyboardMarkup(row_width=1)
    markup.add(
        types.InlineKeyboardButton("⬆️ Subir material",   callback_data="her_subir"),
        types.InlineKeyboardButton("🔍 Buscar material",  callback_data="her_buscar"),
        types.InlineKeyboardButton("⬅️ Volver",           callback_data="her_menu"),
    )
    bot.edit_message_text(
        "📂 *Banco de Material TUCE*\n\n"
        "Compartí y descargá parciales, resúmenes y apuntes de tus compañeros.",
        call.message.chat.id, call.message.message_id,
        reply_markup=markup, parse_mode="Markdown"
    )

def iniciar_subida(bot, call):
    """Paso 1 subida: elegir materia."""
    user_id = call.from_user.id
    estados_herramientas[user_id] = {"paso": "mat_subir"}
    bot.answer_callback_query(call.id)

    markup = types.InlineKeyboardMarkup(row_width=1)
    for mat in MATERIAS_BANCO:
        markup.add(types.InlineKeyboardButton(mat, callback_data=f"her_mat_{mat}"))

    bot.send_message(
        call.message.chat.id,
        "📚 *¿Para qué materia subís el material?*",
        reply_markup=markup, parse_mode="Markdown"
    )

def paso_materia_subir(bot, call, materia):
    """Paso 2: elegir tipo."""
    user_id = call.from_user.id
    if user_id not in estados_herramientas:
        return
    estados_herramientas[user_id]["materia"] = materia
    estados_herramientas[user_id]["paso"] = "tipo_subir"

    markup = types.InlineKeyboardMarkup(row_width=1)
    for tipo in TIPOS_MATERIAL:
        markup.add(types.InlineKeyboardButton(tipo, callback_data=f"her_tipo_{tipo}"))

    bot.edit_message_text(
        f"✅ Materia: *{materia}*\n\n🏷️ *¿Qué tipo de material es?*",
        call.message.chat.id, call.message.message_id,
        reply_markup=markup, parse_mode="Markdown"
    )

def paso_tipo_subir(bot, call, tipo):
    """Paso 3: pedir el archivo."""
    user_id = call.from_user.id
    if user_id not in estados_herramientas:
        return
    estados_herramientas[user_id]["tipo"] = tipo
    estados_herramientas[user_id]["paso"] = "archivo_subir"

    markup_cancel = types.InlineKeyboardMarkup()
    markup_cancel.add(types.InlineKeyboardButton("❌ Cancelar", callback_data="her_banco"))
    bot.edit_message_text(
        f"✅ Tipo: *{tipo}*\n\n"
        f"📎 *Ahora enviá el archivo* (PDF, foto, Word o imagen)\n\n"
        f"_Máximo 20MB_",
        call.message.chat.id, call.message.message_id,
        reply_markup=markup_cancel,
        parse_mode="Markdown"
    )

def paso_archivo_subir(bot, message, firebase_db) -> bool:
    """Paso 4: recibe archivo y guarda en Firebase."""
    user_id = message.from_user.id
    if user_id not in estados_herramientas or estados_herramientas[user_id].get("paso") != "archivo_subir":
        return False

    estado = estados_herramientas[user_id]
    file_id = None
    file_name = "archivo"
    file_type = "documento"

    if message.document:
        file_id   = message.document.file_id
        file_name = message.document.file_name or "archivo"
        file_type = "documento"
    elif message.photo:
        file_id   = message.photo[-1].file_id
        file_name = "foto.jpg"
        file_type = "imagen"
    else:
        bot.send_message(message.chat.id, "⚠️ Enviá un archivo o foto válida.")
        return True

    # Guardar en Firebase
    try:
        registro = {
            "materia":    estado["materia"],
            "tipo":       estado["tipo"],
            "file_id":    file_id,
            "file_name":  file_name,
            "file_type":  file_type,
            "uploader_id": str(user_id),
            "uploader_name": message.from_user.first_name or "Anónimo",
            "timestamp":  time.strftime("%Y-%m-%d %H:%M:%S"),
        }
        firebase_db.reference('banco_material').push(registro)

        del estados_herramientas[user_id]

        markup = types.InlineKeyboardMarkup(row_width=1)
        markup.add(
            types.InlineKeyboardButton("🔍 Ver material de esta materia", callback_data=f"her_ver_{estado['materia']}"),
            types.InlineKeyboardButton("⬅️ Volver al banco",              callback_data="her_banco"),
        )
        bot.send_message(
            message.chat.id,
            f"✅ *¡Material subido!*\n\n"
            f"📚 *Materia:* {estado['materia']}\n"
            f"🏷️ *Tipo:* {estado['tipo']}\n\n"
            f"Tus compañeros ya pueden descargarlo.",
            reply_markup=markup, parse_mode="Markdown"
        )
    except Exception as e:
        print(f"❌ Error subiendo material: {e}")
        bot.send_message(message.chat.id, "⚠️ Error al guardar. Intentá de nuevo.")

    return True

def mostrar_buscar_materia(bot, call):
    """Explorar material por materia."""
    markup = types.InlineKeyboardMarkup(row_width=1)
    for mat in MATERIAS_BANCO:
        markup.add(types.InlineKeyboardButton(mat, callback_data=f"her_ver_{mat}"))
    markup.add(types.InlineKeyboardButton("⬅️ Volver", callback_data="her_banco"))
    bot.edit_message_text(
        "🔍 *¿De qué materia buscás material?*",
        call.message.chat.id, call.message.message_id,
        reply_markup=markup, parse_mode="Markdown"
    )

def mostrar_material_materia(bot, call, materia, firebase_db):
    """Lista el material disponible de una materia."""
    bot.answer_callback_query(call.id, "Buscando material...")
    try:
        ref  = firebase_db.reference('banco_material')
        data = ref.get()
        if not data:
            items = []
        else:
            items = [v for v in data.values() if v.get("materia") == materia]
    except Exception as e:
        print(f"❌ Error leyendo banco: {e}")
        items = []

    if not items:
        markup = types.InlineKeyboardMarkup()
        markup.add(
            types.InlineKeyboardButton("⬆️ Ser el primero en subir", callback_data="her_subir"),
            types.InlineKeyboardButton("⬅️ Volver",                  callback_data="her_buscar"),
        )
        bot.edit_message_text(
            f"😕 Todavía no hay material de *{materia}*.\n¡Podés ser el primero en subir!",
            call.message.chat.id, call.message.message_id,
            reply_markup=markup, parse_mode="Markdown"
        )
        return

    # Ordenar por fecha desc
    items.sort(key=lambda x: x.get("timestamp", ""), reverse=True)

    texto = f"📚 *{materia}*\n{len(items)} archivo(s) disponible(s)\n\n"
    markup = types.InlineKeyboardMarkup(row_width=1)

    for item in items[:15]:
        nombre   = item.get("uploader_name", "Anónimo")
        tipo     = item.get("tipo", "")
        fecha    = item.get("timestamp", "")[:10]
        file_id  = item.get("file_id", "")
        texto += f"{tipo} · {nombre} · {fecha}\n"
        markup.add(types.InlineKeyboardButton(
            f"⬇️ {tipo} — {nombre}",
            callback_data=f"her_dl_{file_id[:40]}"
        ))

    markup.add(types.InlineKeyboardButton("⬅️ Volver", callback_data="her_buscar"))
    bot.edit_message_text(
        texto, call.message.chat.id, call.message.message_id,
        reply_markup=markup, parse_mode="Markdown"
    )

def descargar_archivo(bot, call, file_id_short, firebase_db):
    """Reenvía el archivo al usuario."""
    bot.answer_callback_query(call.id, "Enviando archivo...")
    try:
        ref  = firebase_db.reference('banco_material')
        data = ref.get()
        if not data:
            bot.send_message(call.message.chat.id, "❌ Archivo no encontrado.")
            return
        item = next((v for v in data.values() if v.get("file_id", "").startswith(file_id_short)), None)
        if not item:
            bot.send_message(call.message.chat.id, "❌ Archivo no encontrado.")
            return

        fid  = item["file_id"]
        tipo = item.get("file_type", "documento")
        nombre = item.get("uploader_name", "Anónimo")
        mat    = item.get("materia", "")

        caption = f"📎 *{item.get('tipo','')}*\n📚 {mat}\n👤 Subido por {nombre}"

        if tipo == "imagen":
            bot.send_photo(call.message.chat.id, fid, caption=caption, parse_mode="Markdown")
        else:
            bot.send_document(call.message.chat.id, fid, caption=caption, parse_mode="Markdown")
    except Exception as e:
        print(f"❌ Error descargando: {e}")
        bot.send_message(call.message.chat.id, "⚠️ Error al enviar el archivo. Intentá de nuevo.")


# ─────────────────────────────────────────────
#  FOTO → WORD / PDF  (OCR con GPT-4o)
# ─────────────────────────────────────────────

def menu_ocr(bot, call):
    user_id = call.from_user.id
    estados_herramientas[user_id] = {"paso": "esperando_foto_ocr"}
    bot.answer_callback_query(call.id)
    markup_cancel = types.InlineKeyboardMarkup()
    markup_cancel.add(types.InlineKeyboardButton("❌ Cancelar", callback_data="her_menu"))
    bot.edit_message_text(
        "📸 *Foto → Word / PDF*\n\n"
        "Sacá una foto a tu apunte o al pizarrón y yo lo convierto en un documento.\n\n"
        "📎 *Enviame la foto ahora:*",
        call.message.chat.id, call.message.message_id,
        reply_markup=markup_cancel,
        parse_mode="Markdown"
    )

def procesar_foto_ocr(bot, message, openai_client) -> bool:
    """Recibe foto, llama a GPT-4o Vision y genera Word + PDF."""
    user_id = message.from_user.id
    if user_id not in estados_herramientas or estados_herramientas[user_id].get("paso") != "esperando_foto_ocr":
        return False

    if not message.photo:
        return False

    del estados_herramientas[user_id]
    bot.send_chat_action(message.chat.id, "typing")
    bot.send_message(message.chat.id, "⏳ Procesando tu foto con IA... Un momento.")

    try:
        # Descargar la foto de Telegram
        file_info = bot.get_file(message.photo[-1].file_id)
        file_url  = f"https://api.telegram.org/file/bot{bot.token}/{file_info.file_path}"
        img_resp  = requests.get(file_url, timeout=15)
        img_b64   = base64.b64encode(img_resp.content).decode("utf-8")

        # Llamar a GPT-4o Vision
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Sos un asistente que convierte fotos de apuntes o pizarrones en documentos de texto estructurados. "
                            "Transcribí TODO el texto que ves en la imagen, preservando la estructura: "
                            "títulos, subtítulos, listas, fórmulas y párrafos. "
                            "Formatéalo de forma clara con secciones bien definidas. "
                            "Si hay diagramas o esquemas, descríbelos brevemente entre corchetes [Diagrama: ...]. "
                            "Respondé SOLO con el texto estructurado, sin explicaciones previas."
                        )
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}
                    }
                ]
            }],
            max_tokens=2000
        )

        texto_extraido = response.choices[0].message.content.strip()

        if not texto_extraido:
            bot.send_message(message.chat.id, "⚠️ No pude leer texto en la imagen. Intentá con mejor iluminación.")
            return True

        # Preguntar formato
        estados_herramientas[user_id] = {
            "paso":  "elegir_formato_ocr",
            "texto": texto_extraido
        }

        markup = types.InlineKeyboardMarkup(row_width=2)
        markup.add(
            types.InlineKeyboardButton("📄 Word (.docx)", callback_data="her_fmt_word"),
            types.InlineKeyboardButton("📕 PDF",          callback_data="her_fmt_pdf"),
            types.InlineKeyboardButton("📄+📕 Ambos",     callback_data="her_fmt_ambos"),
        )
        preview = texto_extraido[:300] + ("..." if len(texto_extraido) > 300 else "")
        bot.send_message(
            message.chat.id,
            f"✅ *Texto extraído:*\n\n_{preview}_\n\n¿En qué formato lo querés?",
            reply_markup=markup, parse_mode="Markdown"
        )

    except Exception as e:
        print(f"❌ Error OCR: {e}")
        bot.send_message(message.chat.id, "⚠️ Error al procesar la imagen. Intentá de nuevo.")

    return True

def generar_documento(bot, call, formato: str, openai_client):
    """Genera Word y/o PDF desde el texto extraído."""
    user_id = call.from_user.id
    if user_id not in estados_herramientas or estados_herramientas[user_id].get("paso") != "elegir_formato_ocr":
        bot.answer_callback_query(call.id, "Sesión expirada. Enviá la foto de nuevo.")
        return

    texto = estados_herramientas[user_id]["texto"]
    del estados_herramientas[user_id]

    bot.answer_callback_query(call.id)
    bot.send_chat_action(call.message.chat.id, "upload_document")
    bot.edit_message_text("⏳ Generando documento...", call.message.chat.id, call.message.message_id)

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import fpdf

        lineas = texto.split("\n")

        # ── Generar Word ──
        def crear_word():
            doc = Document()
            doc.core_properties.author = "Bot TUCE - Bytes Creativos"

            # Estilo del documento
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            for linea in lineas:
                linea = linea.strip()
                if not linea:
                    doc.add_paragraph("")
                    continue
                # Detectar títulos (líneas cortas en mayúsculas o con ##)
                if linea.startswith("## "):
                    p = doc.add_heading(linea[3:], level=2)
                elif linea.startswith("# "):
                    p = doc.add_heading(linea[2:], level=1)
                elif linea.isupper() and len(linea) < 80:
                    p = doc.add_heading(linea, level=1)
                elif linea.startswith("- ") or linea.startswith("• "):
                    p = doc.add_paragraph(linea[2:], style='List Bullet')
                elif linea[0].isdigit() and len(linea) > 2 and linea[1] in '.):':
                    p = doc.add_paragraph(linea, style='List Number')
                else:
                    p = doc.add_paragraph(linea)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        # ── Generar PDF ──
        def crear_pdf():
            pdf = fpdf.FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=11)

            for linea in lineas:
                linea = linea.strip()
                if not linea:
                    pdf.ln(4)
                    continue
                if linea.startswith("# ") or (linea.isupper() and len(linea) < 80):
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.multi_cell(0, 8, linea.replace("# ", ""))
                    pdf.set_font("Helvetica", size=11)
                elif linea.startswith("## "):
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.multi_cell(0, 7, linea.replace("## ", ""))
                    pdf.set_font("Helvetica", size=11)
                elif linea.startswith("- ") or linea.startswith("• "):
                    pdf.multi_cell(0, 6, f"  - {linea[2:]}")
                else:
                    pdf.multi_cell(0, 6, linea)

            buf = io.BytesIO()
            pdf.output(buf)
            buf.seek(0)
            return buf

        fecha = time.strftime("%Y%m%d_%H%M")

        if formato in ["word", "ambos"]:
            word_buf = crear_word()
            bot.send_document(
                call.message.chat.id,
                word_buf,
                visible_file_name=f"apunte_tuce_{fecha}.docx",
                caption="📄 *Apunte en Word* — generado por Bot TUCE",
                parse_mode="Markdown"
            )

        if formato in ["pdf", "ambos"]:
            pdf_buf = crear_pdf()
            bot.send_document(
                call.message.chat.id,
                pdf_buf,
                visible_file_name=f"apunte_tuce_{fecha}.pdf",
                caption="📕 *Apunte en PDF* — generado por Bot TUCE",
                parse_mode="Markdown"
            )

        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("🛠️ Usar otra herramienta", callback_data="her_menu_msg"))
        bot.send_message(
            call.message.chat.id,
            "✅ *¡Listo!* Documento generado con éxito.\n\n"
            "_¿Querés convertir otra foto? Solo enviá una nueva imagen._",
            reply_markup=markup, parse_mode="Markdown"
        )

    except Exception as e:
        print(f"❌ Error generando doc: {e}")
        bot.send_message(call.message.chat.id, "⚠️ Error al generar el documento. Intentá de nuevo.")
